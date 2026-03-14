"""PDF operations module for rotating and merging PDF files."""

from pypdf import PdfReader, PdfWriter
from pdf2docx import Converter
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
import os


def merge_pdfs(file1_path: str, file2_path: str, output_path: str) -> None:
    """
    Merges two PDF files into a single PDF.

    Args:
        file1_path (str): Path to the first PDF file
        file2_path (str): Path to the second PDF file
        output_path (str): Path to save the merged PDF file
    """
    writer = PdfWriter()

    # Add all pages from first PDF
    with open(file1_path, 'rb') as pdf1_file:
        reader1 = PdfReader(pdf1_file)
        for page in reader1.pages:
            writer.add_page(page)

    # Add all pages from second PDF
    with open(file2_path, 'rb') as pdf2_file:
        reader2 = PdfReader(pdf2_file)
        for page in reader2.pages:
            writer.add_page(page)

    # Write the merged PDF to output file
    with open(output_path, 'wb') as output_file:
        writer.write(output_file)

    print(f"Successfully merged '{file1_path}' and '{file2_path}' into '{output_path}'")


def rotate_pdf_pages(input_path: str, output_path: str, pages_to_rotate: list[int], rotation_angle: int = 180) -> None:
    """
    Rotates specified pages in a PDF file by a given angle (90, 180, or 270 degrees).

    Args:
        input_path (str): The path to the input PDF file.
        output_path (str): The path to save the output PDF file.
        pages_to_rotate (list of int): A list of page numbers (1-based) to rotate.
        rotation_angle (int): The clockwise rotation angle (90, 180, 270).
    """
    # Open the original PDF file
    with open(input_path, 'rb') as pdf_in_file:
        reader = PdfReader(pdf_in_file)
        writer = PdfWriter()

        # Iterate through all pages
        for page_num in range(len(reader.pages)):
            page = reader.pages[page_num]
            # Check if the current page number is in the list to rotate (convert to 0-based index)
            if (page_num + 1) in pages_to_rotate:
                page.rotate(rotation_angle)
            writer.add_page(page)

        # Write the modified PDF to a new file
        with open(output_path, 'wb') as pdf_out_file:
            writer.write(pdf_out_file)

    print(f"Successfully rotated pages {pages_to_rotate} of '{input_path}' by {rotation_angle} degrees and saved as '{output_path}'")


def order_pdf_pages(input_path: str, output_path: str, num_pages: int, new_order: str) -> None:
    """
    Reorders the first N pages of a PDF file according to the specified order.

    Args:
        input_path (str): The path to the input PDF file.
        output_path (str): The path to save the output PDF file.
        num_pages (int): Number of pages to reorder from the beginning.
        new_order (str): Comma-separated list of page positions (e.g., '1,3,2,5,4,6').
    """
    with open(input_path, 'rb') as pdf_in_file:
        reader = PdfReader(pdf_in_file)
        total_pages = len(reader.pages)

        if num_pages > total_pages:
            raise ValueError(f"Cannot reorder {num_pages} pages: PDF only has {total_pages} pages")

        # Parse the new order
        try:
            order_list = [int(x.strip()) for x in new_order.split(',')]
        except ValueError as e:
            raise ValueError(f"Invalid new_order format: {e}. Expected comma-separated integers.")

        if len(order_list) != num_pages:
            raise ValueError(f"new_order must contain exactly {num_pages} page positions, got {len(order_list)}")

        # Validate that order_list contains each number from 1 to num_pages exactly once
        expected = set(range(1, num_pages + 1))
        actual = set(order_list)
        if expected != actual:
            raise ValueError(f"new_order must contain each number from 1 to {num_pages} exactly once. Got: {order_list}")

        # Create the reordered page list
        reordered_pages = []
        for pos in order_list:
            reordered_pages.append(reader.pages[pos - 1])  # Convert to 0-based index

        # Add any remaining pages in their original order
        for i in range(num_pages, total_pages):
            reordered_pages.append(reader.pages[i])

        # Create new PDF with reordered pages
        writer = PdfWriter()
        for page in reordered_pages:
            writer.add_page(page)

        # Write the reordered PDF to output file
        with open(output_path, 'wb') as pdf_out_file:
            writer.write(pdf_out_file)

    print(f"Successfully reordered first {num_pages} pages of '{input_path}' according to '{new_order}' and saved as '{output_path}'")


def convert_pdf_to_docx(input_path: str, output_path: str) -> None:
    """
    Converts a PDF file to a Word DOCX document.

    Args:
        input_path (str): Path to the input PDF file
        output_path (str): Path to save the output DOCX file
    """
    try:
        # First try the standard pdf2docx conversion
        cv = Converter(input_path)
        cv.convert(output_path, start=0, end=None)
        cv.close()
        print(f"Successfully converted '{input_path}' to DOCX format and saved as '{output_path}'")

    except Exception as e:
        print(f"Standard conversion failed: {e}")
        print("Attempting fallback text extraction method...")

        try:
            # Fallback: Extract text using PyMuPDF and create a basic DOCX
            _fallback_text_extraction(input_path, output_path)
            print(f"Successfully converted '{input_path}' to DOCX format using fallback method and saved as '{output_path}'")

        except Exception as fallback_error:
            print(f"Fallback conversion also failed: {fallback_error}")
            raise RuntimeError(f"PDF conversion failed for '{input_path}'. The PDF may contain complex layouts or unsupported elements.")


def _fallback_text_extraction(input_path: str, output_path: str) -> None:
    """
    Fallback method that extracts text from PDF and creates a basic DOCX document.
    Useful for PDFs with complex layouts that pdf2docx can't handle.

    Args:
        input_path (str): Path to the input PDF file
        output_path (str): Path to save the output DOCX file
    """
    # Create a new Word document
    doc = Document()
    doc.add_heading('PDF Content', 0)

    # Open PDF with PyMuPDF
    pdf_doc = fitz.open(input_path)

    try:
        for page_num in range(len(pdf_doc)):
            page = pdf_doc[page_num]

            # Add page heading
            doc.add_heading(f'Page {page_num + 1}', level=1)

            # Extract text from the page
            text = page.get_text()

            # Split text into paragraphs and add to document
            paragraphs = text.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():  # Only add non-empty paragraphs
                    doc.add_paragraph(para_text.strip())

            # Add page break between pages (except for the last page)
            if page_num < len(pdf_doc) - 1:
                doc.add_page_break()

    finally:
        pdf_doc.close()

    # Save the document
    doc.save(output_path)